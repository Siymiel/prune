"""Knowledge base router — manage knowledge bases and ingest documents into Pinecone."""

from __future__ import annotations

import asyncio
import io
import os
import tempfile
import uuid
from datetime import UTC, datetime
from typing import Any

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Query, Response, UploadFile
from pydantic import BaseModel
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from prune_api.core.auth import CurrentUser, get_current_user
from prune_api.core.settings import settings
from prune_api.db.base import get_session
from prune_api.db.models import KnowledgeBase, KnowledgeDocument

router = APIRouter()

SUPPORTED_TYPES = {"pdf", "docx", "txt", "md"}


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def _extract_text(filename: str, content: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    if ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if ext == "docx":
        from docx import Document
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    return content.decode("utf-8", errors="replace")


def _chunk_text(
    text: str,
    size: int = 500,
    overlap_pct: int = 20,
    method: str = "sentence",
) -> list[str]:
    """Split text into chunks using either naive (fixed-size) or sentence-based method.

    Args:
        size:        target chunk size in characters (naive) or approximate token-equivalent
        overlap_pct: percentage of chunk to repeat in the next chunk (0-50)
        method:      "naive" splits by character count; "sentence" splits at sentence boundaries
    """
    overlap = max(0, int(size * overlap_pct / 100))

    if method == "sentence":
        return _chunk_sentences(text, size, overlap)

    # Naive fixed-length chunking
    chunks: list[str] = []
    step = max(1, size - overlap)
    start = 0
    while start < len(text):
        chunk = text[start : start + size].strip()
        if chunk:
            chunks.append(chunk)
        start += step
    return chunks


def _chunk_sentences(text: str, target_size: int, overlap: int) -> list[str]:
    """Sentence-aware chunker: accumulates sentences until reaching target_size chars."""
    import re
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    sentences = [s.strip() for s in sentences if s.strip()]

    chunks: list[str] = []
    current_sentences: list[str] = []
    current_len = 0

    for sentence in sentences:
        slen = len(sentence)
        if current_len + slen > target_size and current_sentences:
            chunk = " ".join(current_sentences).strip()
            if chunk:
                chunks.append(chunk)
            # Seed next chunk with overlap characters from the end
            overlap_text = chunk[-overlap:] if overlap > 0 else ""
            current_sentences = [overlap_text, sentence] if overlap_text else [sentence]
            current_len = len(overlap_text) + slen + (1 if overlap_text else 0)
        else:
            current_sentences.append(sentence)
            current_len += slen + 1

    if current_sentences:
        chunk = " ".join(current_sentences).strip()
        if chunk:
            chunks.append(chunk)

    return chunks or [text[:target_size].strip()]


# ---------------------------------------------------------------------------
# Background ingestion
# ---------------------------------------------------------------------------

async def _summarize_chunks(chunks: list[str]) -> list[str]:
    """Compress each chunk to a dense summary using Claude Haiku.

    Falls back to the original text if the API key is missing or the call fails,
    so ingestion always succeeds even if summarization is unavailable.
    """
    if not settings.anthropic_api_key:
        return chunks

    from anthropic import AsyncAnthropic
    client = AsyncAnthropic(api_key=settings.anthropic_api_key)
    sem = asyncio.Semaphore(5)

    async def _one(chunk: str) -> str:
        async with sem:
            try:
                msg = await client.messages.create(
                    model="claude-haiku-4-5-20251001",
                    max_tokens=150,
                    system="You are a precise document summarizer. Output ONLY the summary — no preamble, no labels.",
                    messages=[{
                        "role": "user",
                        "content": (
                            "Summarize the following passage in 2-3 dense sentences, "
                            "preserving all key facts, figures, names, and technical details:\n\n"
                            + chunk
                        ),
                    }],
                )
                return msg.content[0].text
            except Exception:
                return chunk  # safe fallback

    return list(await asyncio.gather(*[_one(c) for c in chunks]))


async def _embed_texts(texts: list[str], model: str) -> list[list[float]]:
    """Embed a list of texts using the specified model.

    Supports:
      - Voyage AI models (voyage-3, voyage-3-lite, …)   — requires VOYAGE_API_KEY
      - OpenAI embedding models (text-embedding-3-*)    — requires OPENAI_API_KEY
      - Fallback for unknown models: Voyage AI voyage-3
    """
    model_lower = model.lower()

    if model_lower.startswith("text-embedding") or model_lower in ("all-mpnet-base-v2", "bert-base-cased"):
        # OpenAI-compatible or open-source model via OpenAI-compatible endpoint
        if not settings.openai_api_key:
            raise RuntimeError("OPENAI_API_KEY is not configured for OpenAI embedding models")
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=settings.openai_api_key)
        resp = await client.embeddings.create(input=texts, model=model_lower if model_lower.startswith("text-embedding") else "text-embedding-3-small")
        return [item.embedding for item in resp.data]

    # Voyage AI (default)
    if not settings.voyage_api_key:
        raise RuntimeError("VOYAGE_API_KEY is not configured")
    import voyageai
    vo = voyageai.AsyncClient(api_key=settings.voyage_api_key)
    voyage_model = model if model.startswith("voyage-") else "voyage-3"
    embed_resp = await vo.embed(texts, model=voyage_model, input_type="document")
    return embed_resp.embeddings


async def _ingest(
    doc_id: str,
    kb_id: str,
    filename: str,
    tmp_path: str,
    chunk_size: int = 500,
    chunk_overlap_pct: int = 20,
    chunking_method: str = "sentence",
    embedding_model: str = "voyage-3",
) -> None:
    """Extract text, embed, upsert chunks to Pinecone, update DB status.

    tmp_path is a temporary file written during upload.  It is always deleted
    before this function returns, regardless of success or failure.
    """
    from pinecone import Pinecone
    from prune_api.db.base import AsyncSessionLocal

    try:
        async with AsyncSessionLocal() as session:
            row = await session.get(KnowledgeDocument, uuid.UUID(doc_id))
            if row is None:
                return
            try:
                with open(tmp_path, "rb") as fh:
                    content = fh.read()
                text = _extract_text(filename, content)
                del content  # release upload bytes before the embedding calls

                chunks = _chunk_text(text, size=chunk_size, overlap_pct=chunk_overlap_pct, method=chunking_method)
                if not chunks:
                    row.status = "error"
                    row.error = "No text could be extracted from the file."
                    await session.commit()
                    return

                # Summarize chunks for token-efficient retrieval; keep originals for "show source"
                summaries = await _summarize_chunks(chunks)

                vectors = await _embed_texts(chunks, embedding_model)

                if not settings.pinecone_api_key:
                    raise RuntimeError("PINECONE_API_KEY is not configured")
                pc = Pinecone(api_key=settings.pinecone_api_key)
                index = pc.Index(settings.pinecone_index)

                to_upsert = [
                    {
                        "id": f"{doc_id}-{i}",
                        "values": vectors[i],
                        "metadata": {
                            "text": chunks[i],
                            "summary": summaries[i],
                            "doc_id": doc_id,
                            "kb_id": kb_id,
                            "filename": filename,
                            "chunk_index": i,
                            "embedding_model": embedding_model,
                        },
                    }
                    for i in range(len(chunks))
                ]

                batch = 100
                for start in range(0, len(to_upsert), batch):
                    await asyncio.to_thread(
                        index.upsert,
                        vectors=to_upsert[start : start + batch],
                        namespace=kb_id,
                    )

                row.status = "ready"
                row.chunk_count = len(chunks)
            except Exception as exc:
                row.status = "error"
                row.error = str(exc)

            await session.commit()
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------

class KnowledgeBaseCreate(BaseModel):
    name: str
    description: str | None = None


class KnowledgeBaseOut(BaseModel):
    id: str
    name: str
    description: str | None
    created_at: str
    document_count: int = 0


class KnowledgeDocumentOut(BaseModel):
    id: str
    filename: str
    file_type: str
    chunk_count: int
    status: str
    error: str | None
    created_at: str


# ---------------------------------------------------------------------------
# Knowledge base CRUD
# ---------------------------------------------------------------------------

@router.get("/knowledge-bases", response_model=list[KnowledgeBaseOut])
async def list_knowledge_bases(
    response: Response,
    limit: int = Query(default=50, ge=1, le=200),
    offset: int = Query(default=0, ge=0),
    current_user: CurrentUser = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> list[KnowledgeBaseOut]:
    total = (await session.scalar(
        select(func.count()).select_from(KnowledgeBase)
        .where(KnowledgeBase.tenant_id == current_user.tenant_id)
    )) or 0
    response.headers["X-Total-Count"] = str(total)

    rows = await session.execute(
        select(KnowledgeBase)
        .where(KnowledgeBase.tenant_id == current_user.tenant_id)
        .order_by(KnowledgeBase.created_at.desc())
        .limit(limit)
        .offset(offset)
    )
    kbs = rows.scalars().all()
    result = []
    for kb in kbs:
        doc_count = (await session.scalar(
            select(func.count()).select_from(KnowledgeDocument)
            .where(KnowledgeDocument.knowledge_base_id == kb.id)
        )) or 0
        result.append(KnowledgeBaseOut(
            id=str(kb.id),
            name=kb.name,
            description=kb.description,
            created_at=kb.created_at.isoformat(),
            document_count=doc_count,
        ))
    return result


@router.post("/knowledge-bases", response_model=KnowledgeBaseOut, status_code=201)
async def create_knowledge_base(
    body: KnowledgeBaseCreate,
    current_user: CurrentUser = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> KnowledgeBaseOut:
    kb = KnowledgeBase(
        tenant_id=current_user.tenant_id,
        name=body.name,
        description=body.description,
    )
    session.add(kb)
    await session.commit()
    await session.refresh(kb)
    return KnowledgeBaseOut(
        id=str(kb.id),
        name=kb.name,
        description=kb.description,
        created_at=kb.created_at.isoformat(),
        document_count=0,
    )


@router.delete("/knowledge-bases/{kb_id}", status_code=204)
async def delete_knowledge_base(
    kb_id: str,
    current_user: CurrentUser = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> None:
    try:
        kid = uuid.UUID(kb_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid kb_id")

    row = await session.execute(
        select(KnowledgeBase).where(
            KnowledgeBase.id == kid,
            KnowledgeBase.tenant_id == current_user.tenant_id,
        )
    )
    kb = row.scalar_one_or_none()
    if kb is None:
        raise HTTPException(status_code=404, detail="Knowledge base not found")

    # Delete all vectors from Pinecone namespace; abort if cleanup fails so no orphans are left
    if settings.pinecone_api_key:
        try:
            from pinecone import Pinecone
            pc = Pinecone(api_key=settings.pinecone_api_key)
            index = pc.Index(settings.pinecone_index)
            await asyncio.to_thread(index.delete, delete_all=True, namespace=kb_id)
        except Exception as exc:
            raise HTTPException(
                status_code=500,
                detail=f"Pinecone cleanup failed — knowledge base not deleted: {exc}",
            )

    await session.delete(kb)
    await session.commit()


# ---------------------------------------------------------------------------
# Document upload + listing
# ---------------------------------------------------------------------------

@router.get("/knowledge-bases/{kb_id}/documents", response_model=list[KnowledgeDocumentOut])
async def list_documents(
    kb_id: str,
    response: Response,
    limit: int = Query(default=50, ge=1, le=200),
    offset: int = Query(default=0, ge=0),
    current_user: CurrentUser = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> list[KnowledgeDocumentOut]:
    try:
        kid = uuid.UUID(kb_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid kb_id")

    kb_row = await session.execute(
        select(KnowledgeBase).where(
            KnowledgeBase.id == kid,
            KnowledgeBase.tenant_id == current_user.tenant_id,
        )
    )
    if kb_row.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Knowledge base not found")

    total = (await session.scalar(
        select(func.count()).select_from(KnowledgeDocument)
        .where(KnowledgeDocument.knowledge_base_id == kid)
    )) or 0
    response.headers["X-Total-Count"] = str(total)

    rows = await session.execute(
        select(KnowledgeDocument)
        .where(KnowledgeDocument.knowledge_base_id == kid)
        .order_by(KnowledgeDocument.created_at)
        .limit(limit)
        .offset(offset)
    )
    return [
        KnowledgeDocumentOut(
            id=str(d.id),
            filename=d.filename,
            file_type=d.file_type,
            chunk_count=d.chunk_count,
            status=d.status,
            error=d.error,
            created_at=d.created_at.isoformat(),
        )
        for d in rows.scalars().all()
    ]


@router.post("/knowledge-bases/{kb_id}/documents", response_model=KnowledgeDocumentOut, status_code=202)
async def upload_document(
    kb_id: str,
    file: UploadFile,
    background_tasks: BackgroundTasks,
    chunk_size: int = Query(default=500, ge=100, le=4000, description="Chunk size in characters"),
    chunk_overlap_pct: int = Query(default=20, ge=0, le=50, description="Overlap percentage between chunks"),
    chunking_method: str = Query(default="sentence", description="'sentence' or 'naive'"),
    embedding_model: str = Query(default="voyage-3", description="Embedding model to use"),
    current_user: CurrentUser = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> KnowledgeDocumentOut:
    try:
        kid = uuid.UUID(kb_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid kb_id")

    kb_row = await session.execute(
        select(KnowledgeBase).where(
            KnowledgeBase.id == kid,
            KnowledgeBase.tenant_id == current_user.tenant_id,
        )
    )
    if kb_row.scalar_one_or_none() is None:
        raise HTTPException(status_code=404, detail="Knowledge base not found")

    filename = file.filename or "upload"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    if ext not in SUPPORTED_TYPES:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_TYPES))}",
        )

    # Write upload to a temp file so the request handler doesn't hold large bytes
    # in memory for the duration of the background ingestion task.
    fd, tmp_path = tempfile.mkstemp(suffix=f".{ext}")
    try:
        with os.fdopen(fd, "wb") as fh:
            fh.write(await file.read())
    except Exception:
        os.unlink(tmp_path)
        raise

    doc = KnowledgeDocument(
        knowledge_base_id=kid,
        tenant_id=current_user.tenant_id,
        filename=filename,
        file_type=ext,
        status="processing",
    )
    session.add(doc)
    await session.commit()
    await session.refresh(doc)

    background_tasks.add_task(
        _ingest,
        str(doc.id),
        kb_id,
        filename,
        tmp_path,
        chunk_size,
        chunk_overlap_pct,
        chunking_method,
        embedding_model,
    )

    return KnowledgeDocumentOut(
        id=str(doc.id),
        filename=doc.filename,
        file_type=doc.file_type,
        chunk_count=0,
        status="processing",
        error=None,
        created_at=doc.created_at.isoformat(),
    )


@router.delete("/knowledge-bases/{kb_id}/documents/{doc_id}", status_code=204)
async def delete_document(
    kb_id: str,
    doc_id: str,
    current_user: CurrentUser = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> None:
    try:
        kid = uuid.UUID(kb_id)
        did = uuid.UUID(doc_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid id")

    row = await session.execute(
        select(KnowledgeDocument).where(
            KnowledgeDocument.id == did,
            KnowledgeDocument.knowledge_base_id == kid,
            KnowledgeDocument.tenant_id == current_user.tenant_id,
        )
    )
    doc = row.scalar_one_or_none()
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")

    # Remove vectors from Pinecone — only if ingestion succeeded (chunk_count > 0)
    if settings.pinecone_api_key and doc.chunk_count > 0:
        try:
            from pinecone import Pinecone
            pc = Pinecone(api_key=settings.pinecone_api_key)
            index = pc.Index(settings.pinecone_index)
            chunk_ids = [f"{doc_id}-{i}" for i in range(doc.chunk_count)]
            await asyncio.to_thread(index.delete, ids=chunk_ids, namespace=kb_id)
        except Exception:
            pass  # best-effort: DB record still deleted, stale vectors are inert

    await session.delete(doc)
    await session.commit()
